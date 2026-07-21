from types import SimpleNamespace

from docx import Document
from pypdf import PdfWriter
from pypdf.generic import (
    DecodedStreamObject,
    DictionaryObject,
    NameObject,
)

from app.documents.high_fidelity_parser import HighFidelityParser
from app.documents.parser_contract import ParseRequest, ParseResult


class Fallback:
    def __init__(self):
        self.calls = 0

    def parse(self, _request):
        self.calls += 1
        return ParseResult("fallback", "current-loaders", "1")


def test_markitdown_is_used_locally_for_structured_formats(tmp_path):
    path = tmp_path / "manual.pdf"
    path.write_bytes(b"local fixture")

    class Converter:
        def convert(self, source):
            assert source == path
            return SimpleNamespace(markdown="# Structured medical document")

    fallback = Fallback()
    result = HighFidelityParser(Converter(), fallback).parse(
        ParseRequest("job-1", "version-1", path, "pdf")
    )
    assert result.parser_name == "markitdown"
    assert fallback.calls == 0


def test_markitdown_failure_uses_existing_loader_adapter(tmp_path):
    path = tmp_path / "manual.docx"
    path.write_bytes(b"local fixture")

    class Converter:
        def convert(self, _source):
            raise RuntimeError("conversion failed")

    fallback = Fallback()
    result = HighFidelityParser(Converter(), fallback).parse(
        ParseRequest("job-1", "version-1", path, "docx")
    )
    assert result.parser_name == "current-loaders"
    assert result.warnings == ("high_fidelity_fallback",)


def test_real_markitdown_docx_conversion(tmp_path):
    path = tmp_path / "manual.docx"
    document = Document()
    document.add_heading("Aspirin Guidance", level=1)
    document.add_paragraph("Take only the reviewed clinical dose.")
    document.save(path)
    result = HighFidelityParser().parse(
        ParseRequest("job-1", "version-1", path, "docx")
    )
    assert result.parser_name == "markitdown"
    assert "Aspirin Guidance" in result.text


def test_real_markitdown_pdf_conversion(tmp_path):
    path = tmp_path / "manual.pdf"
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    page[NameObject("/Resources")] = DictionaryObject(
        {
            NameObject("/Font"): DictionaryObject(
                {NameObject("/F1"): writer._add_object(font)}
            )
        }
    )
    stream = DecodedStreamObject()
    stream.set_data(b"BT /F1 12 Tf 72 720 Td (Aspirin clinical guidance) Tj ET")
    page[NameObject("/Contents")] = writer._add_object(stream)
    with path.open("wb") as output:
        writer.write(output)
    result = HighFidelityParser().parse(
        ParseRequest("job-1", "version-1", path, "pdf")
    )
    assert result.parser_name == "markitdown"
    assert "Aspirin clinical guidance" in result.text
