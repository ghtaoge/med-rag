"""Word 加载器。python-docx 段落 + 表格提取。"""

from __future__ import annotations

from pathlib import Path

from app.documents.loader.base import DocumentLoader


class WordLoader(DocumentLoader):
    """Word(.docx) 加载器。逐段提取文本，表格转文本。"""

    def load(self, file_path: Path) -> str:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("Word 加载需要 python-docx，请执行: pip install python-docx")

        doc = Document(str(file_path))
        parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        for table in doc.tables:
            table_text = self._table_to_text(table)
            if table_text.strip():
                parts.append(table_text)

        return "\n\n".join(parts)

    def _table_to_text(self, table) -> str:
        """表格转文本（每行用 | 分隔）。"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        return "\n".join(rows)
